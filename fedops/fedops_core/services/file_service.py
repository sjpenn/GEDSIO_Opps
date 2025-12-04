import os
import shutil
from datetime import datetime
from typing import List, Optional
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from fedops_core.db.models import StoredFile, Opportunity
from fedops_core.settings import settings
from fedops_core.services.ai_service import AIService
from fedops_core.prompts import determine_document_type
import pandas as pd
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import pypdfium2 as pdfium

# Ensure upload directory exists
os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

class FileService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.ai_service = AIService()

    async def upload_file(self, file: UploadFile, opportunity_id: Optional[int] = None) -> StoredFile:
        file_path = os.path.join(settings.UPLOAD_DIR, file.filename)
        
        # Save file to disk
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # Get file size
        file_size = os.path.getsize(file_path)
        file_type = file.filename.split('.')[-1].lower() if '.' in file.filename else None

        # Create DB record
        db_file = StoredFile(
            filename=file.filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            opportunity_id=opportunity_id
        )
        self.db.add(db_file)
        await self.db.commit()
        await self.db.refresh(db_file)
        return db_file

    async def get_files(self, opportunity_id: Optional[int] = None) -> List[StoredFile]:
        query = select(StoredFile)
        if opportunity_id:
            query = query.where(StoredFile.opportunity_id == opportunity_id)
        result = await self.db.execute(query)
        return result.scalars().all()

    async def get_file(self, file_id: int) -> Optional[StoredFile]:
        result = await self.db.execute(select(StoredFile).where(StoredFile.id == file_id))
        return result.scalar_one_or_none()

    async def process_file(self, file_id: int):
        db_file = await self.get_file(file_id)
        if not db_file:
            raise ValueError("File not found")

        # 1. Parse Content with AdvancedParser
        parsed_elements = None
        content = None
        
        try:
            from fedops_core.services.advanced_parser import AdvancedParser
            parser = AdvancedParser()
            parsed_elements = parser.parse_document(db_file.file_path)
            # Also get plain text for legacy AI summary
            content = " ".join([el.text for el in parsed_elements])
            db_file.parsed_content = content
            print(f"Parsed {db_file.filename} using AdvancedParser: {len(parsed_elements)} elements")
        except Exception as e:
            print(f"AdvancedParser failed for {db_file.filename}: {e}, falling back to legacy parser")
            # Fallback to legacy parsing
            content = self._parse_file_content(db_file.file_path, db_file.file_type)
            db_file.parsed_content = content

        # 2. Generate Summary (Shipley)
        if content:
            try:
                doc_type = determine_document_type(db_file.filename, content)
                response_text = await self.ai_service.generate_shipley_summary(content, doc_type)
                
                import json
                import re

                try:
                    # Clean response if it's wrapped in markdown code blocks
                    cleaned_text = re.sub(r'^```json\s*|\s*```$', '', response_text.strip(), flags=re.MULTILINE)
                    data = json.loads(cleaned_text)
                    
                    summary = data.get("markdown_report", "No report generated.")
                    structured_data = data.get("structured_data", {})
                    
                    db_file.content_summary = summary
                    db_file.analysis_json = structured_data
                except json.JSONDecodeError:
                    # Fallback if response is not valid JSON (e.g. raw text)
                    db_file.content_summary = response_text
                    db_file.analysis_json = {"error": "Failed to parse structured data", "raw_response": response_text}
            except Exception as e:
                print(f"Warning: AI summary generation failed for {db_file.filename}: {e}")
                db_file.content_summary = "AI Summary generation failed."
                db_file.analysis_json = {"error": str(e)}
        
        # 3. Shred Document (Advanced or Legacy)
        try:
            from fedops_core.services.shredding_service import ShreddingService
            shredding_service = ShreddingService(self.db)
            
            if parsed_elements:
                # Use advanced shredding with metadata
                chunk_count = await shredding_service.shred_document(db_file.id, parsed_elements)
                print(f"Advanced shredding: {db_file.filename} into {chunk_count} chunks with page/section metadata")
            elif content:
                # Use legacy shredding
                chunk_count = await shredding_service.shred_document(db_file.id, content)
                print(f"Legacy shredding: {db_file.filename} into {chunk_count} chunks")
        except Exception as e:
            print(f"Warning: Document shredding failed for {db_file.filename}: {e}")

        await self.db.commit()
        await self.db.refresh(db_file)
        return db_file

    async def import_opportunity_resources(self, opportunity_id: int) -> List[StoredFile]:
        # Get opportunity
        stmt = select(Opportunity).where(Opportunity.id == opportunity_id)
        result = await self.db.execute(stmt)
        opportunity = result.scalar_one_or_none()
        
        if not opportunity:
            raise ValueError("Opportunity not found")
            
        imported_files = []
        
        # Use resource_files (resolved) or resource_links (raw)
        resources = opportunity.resource_files or []
        if not resources and opportunity.resource_links:
            resources = [{"url": link, "filename": link.split('/')[-1]} for link in opportunity.resource_links]
            
        import httpx
        async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
            for res in resources:
                url = res.get("url")
                filename = res.get("filename") or "unknown_file"
                
                # Check if already exists for this opportunity
                stmt = select(StoredFile).where(
                    StoredFile.opportunity_id == opportunity_id,
                    StoredFile.filename == filename
                )
                existing = await self.db.execute(stmt)
                existing_file = existing.scalar_one_or_none()
                
                # If file exists and has parsed content, skip
                if existing_file:
                    if existing_file.parsed_content:
                        print(f"File {filename} already imported and processed, skipping")
                        continue
                    else:
                        # File exists but not processed, we'll process it below
                        print(f"File {filename} exists but not processed, will process it")
                        imported_files.append(existing_file)
                        continue
                    
                try:
                    # Download file
                    print(f"Downloading {filename} from {url}")
                    response = await client.get(url)
                    if response.status_code == 200:
                        file_path = os.path.join(settings.UPLOAD_DIR, filename)
                        with open(file_path, "wb") as f:
                            f.write(response.content)
                            
                        file_size = os.path.getsize(file_path)
                        file_type = filename.split('.')[-1].lower() if '.' in filename else None
                        
                        db_file = StoredFile(
                            filename=filename,
                            file_path=file_path,
                            file_type=file_type,
                            file_size=file_size,
                            opportunity_id=opportunity_id
                        )
                        self.db.add(db_file)
                        await self.db.flush()  # Flush to get the ID
                        imported_files.append(db_file)
                        print(f"Successfully downloaded {filename}")
                    else:
                        print(f"Failed to download {url}: HTTP {response.status_code}")
                except Exception as e:
                    print(f"Failed to download {url}: {e}")
                    
        await self.db.commit()
        
        # Now process all imported files to populate parsed_content
        processed_count = 0
        for db_file in imported_files:
            try:
                # Refresh to ensure we have the latest data
                await self.db.refresh(db_file)
                
                # Only process if not already processed
                if not db_file.parsed_content:
                    print(f"Processing file: {db_file.filename}")
                    await self.process_file(db_file.id)
                    processed_count += 1
                    print(f"Successfully processed {db_file.filename}")
            except Exception as e:
                print(f"Warning: Failed to process file {db_file.filename}: {e}")
                # Continue processing other files even if one fails
                
        print(f"Import complete: {len(imported_files)} files imported, {processed_count} files processed")
        return imported_files

    def _parse_file_content(self, file_path: str, file_type: str) -> str:
        try:
            if file_type == 'pdf':
                text = ""
                try:
                    # Method 1: Try pdfplumber (best for layout preservation)
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                            
                            # Extract tables (basic)
                            tables = page.extract_tables()
                            for table in tables:
                                for row in table:
                                    text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                except Exception as e:
                    print(f"pdfplumber failed: {e}")

                # Method 2: Fallback to OCR if text is empty or very short
                if len(text.strip()) < 100:
                    print(f"PDF text extraction yielded {len(text)} chars. Attempting OCR with pypdfium2/tesseract...")
                    try:
                        pdf = pdfium.PdfDocument(file_path)
                        ocr_text = ""
                        # Limit to first 20 pages to avoid timeout on large docs during synchronous processing
                        # In a real async worker we could do all, but for now let's be safe
                        pages_to_process = min(len(pdf), 20)
                        
                        for i in range(pages_to_process):
                            page = pdf[i]
                            # Render to image (scale=2 for better OCR)
                            bitmap = page.render(scale=2)
                            pil_image = bitmap.to_pil()
                            
                            # OCR
                            page_ocr = pytesseract.image_to_string(pil_image)
                            ocr_text += page_ocr + "\n"
                        
                        if len(ocr_text.strip()) > len(text.strip()):
                            text = ocr_text
                            print(f"OCR extracted {len(text)} chars")
                        else:
                            print("OCR did not extract more text than pdfplumber")
                            
                    except Exception as e:
                        print(f"OCR fallback failed: {e}")
                
                return text
            elif file_type in ['xlsx', 'xls']:
                df = pd.read_excel(file_path)
                return df.to_string()
            elif file_type in ['docx', 'doc']:
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            elif file_type in ['jpg', 'jpeg', 'png']:
                try:
                    return pytesseract.image_to_string(Image.open(file_path))
                except:
                    return "[Image content - OCR not available]"
            else:
                with open(file_path, 'r', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            return f"Error parsing file: {str(e)}"
