from docx import Document

doc = Document()
doc.add_heading('Request for Proposal (RFP)', 0)
doc.add_paragraph('This is a test RFP document for Past Performance generation verification.')

doc.add_heading('Scope of Work', level=1)
doc.add_paragraph('The contractor shall provide cybersecurity support services to the Department of Homeland Security.')
doc.add_paragraph('Tasks include: Vulnerability assessment, Threat monitoring, and Incident response.')

doc.add_heading('Technical Approach', level=1)
doc.add_paragraph('We will use Agile methodology and NIST 800-53 standards.')

doc.add_heading('Results', level=1)
doc.add_paragraph('Reduced security incidents by 40% and improved response time by 50%.')

doc.save('test_pp.docx')
